#!/usr/bin/env python3
"""
Resume Scanner — TechyEra Interview Assistant
Scans the Downloads folder for resumes, extracts candidate info, and writes to Excel.
"""

import os
import re
import sys
import subprocess
from pathlib import Path
from datetime import datetime

import fitz                          # PyMuPDF
from docx import Document            # python-docx
import openpyxl
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter

# ── Configuration ─────────────────────────────────────────────────────────────

DOWNLOADS = Path.home() / "Downloads"
RESUME_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

# Keywords that strongly suggest a file is a resume (checked in extracted text)
RESUME_SIGNALS = [
    "experience", "education", "skills", "objective", "summary",
    "work experience", "employment", "qualification", "project",
    "certification", "achievements", "languages", "references",
    "curriculum vitae", "resume", "profile",
]

# Keywords that suggest a file is NOT a resume
NON_RESUME_SIGNALS = [
    "invoice", "receipt", "bank statement", "account statement",
    "offer letter", "appointment letter", "salary slip", "pay slip",
    "insurance", "policy document", "guideline", "rbi", "circular",
    "admit card", "hall ticket", "mark sheet", "scorecard",
    "delay insights", "flight delay", "analysis report", "data analysis",
    "case study", "business analysis", "market research",
]

# Filename substrings that indicate a job description, not a resume
JOB_DESC_SIGNALS = [
    "job description", "job desc", " jd ", "_jd_", "-jd-", "_jd.", "-jd.",
    "job role", "job profile", "position description", "hiring",
    "west coast univ", "job requirement",
]

# ── Text extraction ────────────────────────────────────────────────────────────

def extract_text_pdf(path: Path) -> str:
    try:
        doc = fitz.open(str(path))
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        print(f"    [WARN] PDF read failed: {e}")
        return ""


def extract_text_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs]
        # Also grab table cells — resumes often put info in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paras.append(cell.text)
        return "\n".join(paras)
    except Exception as e:
        print(f"    [WARN] DOCX read failed: {e}")
        return ""


def extract_text_doc(path: Path) -> str:
    # macOS ships textutil; fall back to antiword if available
    for cmd in [
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        ["antiword", str(path)],
    ]:
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
            if result.stdout.strip():
                return result.stdout
        except Exception:
            continue
    return ""


def extract_text_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_pdf(path)
    if ext == ".docx":
        return extract_text_docx(path)
    if ext == ".doc":
        return extract_text_doc(path)
    if ext == ".txt":
        return extract_text_txt(path)
    return ""


# ── Resume detection ───────────────────────────────────────────────────────────

def is_likely_resume(text: str, filename: str) -> bool:
    """Return True if the file looks like a resume, not a bank doc / invoice etc."""
    if not text.strip():
        return False

    fn_lower = filename.lower()
    lower = text.lower()

    # Skip temp / lock files
    if filename.startswith(".~") or filename.startswith("~$"):
        return False

    # Skip obvious job descriptions by filename
    if any(sig in fn_lower for sig in JOB_DESC_SIGNALS):
        return False

    # Hard exclude by content
    if any(sig in lower[:3000] for sig in NON_RESUME_SIGNALS):
        # Still allow if the filename itself says resume/cv
        if not any(w in fn_lower for w in ["resume", "cv", " cv", "_cv"]):
            return False

    # Must have at least 1 resume signal in first 3000 chars
    hits = sum(1 for sig in RESUME_SIGNALS if sig in lower[:3000])
    return hits >= 1


# ── Field extractors ───────────────────────────────────────────────────────────

def extract_emails(text: str) -> list[str]:
    raw = re.findall(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    # Remove placeholder / junk emails
    junk = {"example.com", "email.com", "yourname", "domain.com", "gmail.co", "test.com"}
    seen, out = set(), []
    for e in raw:
        e_low = e.lower()
        if not any(j in e_low for j in junk) and e_low not in seen:
            seen.add(e_low)
            out.append(e)
    return out


def extract_phones(text: str) -> list[str]:
    # Patterns ordered from most-specific to least
    patterns = [
        r"\+91[\s\-]?[6-9]\d{9}",           # +91 Indian
        r"\+1[\s\-]?\(?\d{3}\)?[\s\-]\d{3}[\s\-]\d{4}",  # +1 US
        r"\b[6-9]\d{9}\b",                   # Indian 10-digit
        r"\b\d{3}[\s.\-]\d{3}[\s.\-]\d{4}\b",  # XXX-XXX-XXXX
        r"\+?\d{1,3}[\s\-]\d{4,5}[\s\-]\d{4,6}",  # Generic international
    ]
    raw = []
    for p in patterns:
        raw.extend(re.findall(p, text))

    seen_digits, out = set(), []
    for ph in raw:
        digits = re.sub(r"\D", "", ph)
        if 10 <= len(digits) <= 13 and digits not in seen_digits:
            seen_digits.add(digits)
            out.append(ph.strip())
    return out


def extract_linkedin(text: str) -> str:
    m = re.search(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9\-_%]+)",
        text, re.IGNORECASE
    )
    return f"linkedin.com/in/{m.group(1).rstrip('/')}" if m else ""


def _name_from_filename(filename: str) -> str:
    """
    Heuristically pull a candidate name out of common Indian IT resume
    filename patterns, e.g.:
      "(TCS) Sandeep Tirukovela - Azure Data Engineer.docx"
      "1686707248365_Yarl Balaji Rao.docx"
      "17265_Madhu_Sudhan_Reddy_Yasa_13349_Resume.pdf"
    """
    stem = Path(filename).stem
    stem = re.sub(r"[\(\)\[\]]", " ", stem)    # remove brackets

    NOISE = {
        "resume", "cv", "updated", "update", "new", "final", "latest",
        "onshore", "offshore", "tcs", "wipro", "infosys", "accenture",
        "cognizant", "hcl", "profile", "document",
    }

    # Role/tech words to exclude from name tokens
    ROLE_EXCL = {
        "engineer", "developer", "analyst", "architect", "manager", "data",
        "software", "cloud", "devops", "python", "java", "aws", "azure",
        "dataengineer", "pyth", "cap", "tech", "sr", "jr", "lead", "senior",
    }

    # Strategy A: underscore-separated — filter out pure-numeric and noise tokens
    if "_" in stem:
        tokens = stem.split("_")
        name_tokens = []
        for t in tokens:
            t = t.strip().rstrip(".")      # strip trailing dots (e.g. "vasanth.m" → "vasanth")
            t = t.split(".")[0]            # take part before dot for "vasanth.m"
            if not t or len(t) < 2:
                continue
            if re.fullmatch(r"\d+", t):          # skip pure numbers
                continue
            if re.fullmatch(r"v?\d+", t, re.I):  # skip version tokens v1, v14 …
                continue
            if t.lower() in NOISE or t.lower() in ROLE_EXCL:
                if name_tokens:              # stop if we already started collecting
                    break
                continue
            if re.match(r"^[A-Za-z][a-z]{1,}$", t):
                name_tokens.append(t.capitalize())
                if len(name_tokens) == 4:
                    break
            elif name_tokens:                     # stop at non-name token
                break
        if len(name_tokens) >= 2:
            return " ".join(name_tokens)

    # Strategy B: space/dash separated — cut at dash (role usually follows name)
    stem_clean = re.sub(r"^\d[\d\s\-._]*", "", stem).strip()   # drop leading id/number
    for word in NOISE:
        stem_clean = re.sub(rf"(?i)\b{word}\b", " ", stem_clean)
    stem_clean = re.sub(r"[-|]+.*$", "", stem_clean)            # cut at dash
    stem_clean = re.sub(r"\s+", " ", stem_clean).strip()
    if re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}$", stem_clean):
        return stem_clean

    return ""


# Words that look like names but are actually roles/titles/places — reject these as names
ROLE_WORDS = {
    "engineer", "developer", "analyst", "architect", "manager", "consultant",
    "designer", "lead", "scientist", "specialist", "director", "officer",
    "associate", "intern", "trainee", "fresher", "senior", "junior",
    "software", "data", "cloud", "devops", "fullstack", "frontend", "backend",
    "curriculum", "vitae", "resume", "profile", "summary", "objective",
    "experience", "education", "skills", "contact", "address",
    "professional", "technical", "information", "technology",
    # geographic/org terms that can be mistaken for names
    "united", "kingdom", "states", "america", "india", "google", "amazon",
    "microsoft", "oracle", "apple", "skilled", "experienced", "certified",
    "bachelor", "master", "doctor", "computer", "science", "engineering",
    "subject", "relieving", "letter", "dear", "sincerely",
}


def _looks_like_role(name: str) -> bool:
    words = name.lower().split()
    return any(w in ROLE_WORDS for w in words)


def extract_name(text: str, filename: str) -> str:
    # Priority 1: explicit "Name:" label in document
    m = re.search(
        r"(?i)(?:^|\n)\s*(?:full\s+)?name\s*[:\-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})",
        text
    )
    if m:
        candidate = m.group(1).strip()
        if not _looks_like_role(candidate):
            return candidate

    # Priority 2: filename heuristic
    fn_name = _name_from_filename(filename)
    if fn_name and not _looks_like_role(fn_name):
        return fn_name

    # Priority 3: first 12 non-empty lines — look for a standalone name line
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:12]:
        # Title-case words only, 2-4 tokens, no digits, no special chars
        if re.match(r"^[A-Z][a-z]+(?:[\s.]+[A-Z][a-z]+){1,3}$", line):
            if not _looks_like_role(line) and len(line) < 50:
                return line

    # Priority 4: look for name patterns in first 800 chars
    # Extra geographic / organisation words to exclude from name candidates
    NON_NAME = {
        "united", "kingdom", "states", "america", "india", "google", "amazon",
        "microsoft", "oracle", "apple", "skilled", "experienced", "certified",
        "bachelor", "master", "doctor", "computer", "science", "engineering",
        "hyderabad", "bangalore", "chennai", "mumbai", "pune", "delhi",
    }
    m2 = re.search(
        r"\b([A-Z][a-z]{2,}\s+(?:[A-Z][a-z]{2,}\s+){0,2}[A-Z][a-z]{2,})\b",
        text[:800]
    )
    if m2:
        candidate = m2.group(1).strip()
        words_lower = {w.lower() for w in candidate.split()}
        if (
            not _looks_like_role(candidate)
            and not words_lower & NON_NAME
            and len(candidate.split()) >= 2
        ):
            return candidate

    return "Unknown"


def extract_current_role(text: str, filename: str) -> str:
    # Try to pull from second line of resume (often job title)
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    # Look for explicit label
    for pat in [
        r"(?i)(?:current\s+)?(?:designation|title|position|role)\s*[:\-]\s*(.+)",
        r"(?i)(?:currently\s+working\s+as|working\s+as)\s*[:\-]?\s*(.+)",
    ]:
        m = re.search(pat, text[:3000])
        if m:
            val = m.group(1).strip().split("\n")[0][:80]
            val = re.sub(r"^(?:a|an)\s+", "", val, flags=re.I)   # strip "a/an" prefix
            if val:
                return val

    # Try filename — part after the dash often contains role
    stem = Path(filename).stem
    m = re.search(r"[-–]\s*([A-Za-z ]+(?:Engineer|Developer|Analyst|Architect|Manager|Consultant|Designer|Lead|Scientist|Specialist)[A-Za-z ]*)", stem)
    if m:
        return m.group(1).strip()

    # Second non-name line often has the title (avoid section headers)
    SECTION_HEADERS = {
        "professional summary", "summary", "objective", "career objective",
        "about me", "profile", "overview", "introduction", "education",
        "experience", "work experience", "skills", "projects", "certifications",
        "achievements", "contact", "references", "languages",
    }
    name = extract_name(text, filename)
    for line in lines[1:8]:
        if line.lower() in SECTION_HEADERS:
            continue
        if re.fullmatch(r"[A-Z ]+", line):       # ALL CAPS section header
            continue
        if (
            line != name
            and 4 < len(line) < 70
            and len(line.split()) <= 7          # titles are short, sentences are not
            and "," not in line                 # sentences have commas; titles don't
            and not re.search(r"@|\d{10}|linkedin|github|http|emp\s*id|^\d|\(", line, re.I)
            and not line.lower().startswith(("a ", "an ", "the ", "i am", "i'm", "mr.", "ms.", "subject"))
        ):
            return line

    return ""


def extract_location(text: str) -> str:
    patterns = [
        r"(?i)(?:location|city|address|based\s+(?:in|at))\s*[:\-]\s*([^\n,|]{3,40})",
        r"(?i)(?:hyderabad|bangalore|bengaluru|chennai|mumbai|pune|delhi|noida|"
        r"gurgaon|gurugram|kolkata|ahmedabad|kochi|trivandrum|vizag|coimbatore"
        r"|new york|new jersey|texas|california|chicago|seattle|toronto|london|"
        r"singapore|dubai|sydney)\b",
    ]
    for pat in patterns:
        m = re.search(pat, text[:4000])
        if m:
            return m.group(0 if m.lastindex is None else 1).strip()[:50]
    return ""


def extract_experience_years(text: str) -> str:
    patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*years?\s*(?:of\s*)?(?:overall\s*)?(?:total\s*)?experience",
        r"(?:total|overall)\s*(?:experience\s*(?:of\s*)?)?(\d+(?:\.\d+)?)\+?\s*years?",
        r"experience\s*[:\-]\s*(\d+(?:\.\d+)?)\+?\s*years?",
        r"(\d+(?:\.\d+)?)\+?\s*yrs?\s*(?:of\s*)?(?:exp|experience)",
    ]
    for pat in patterns:
        m = re.search(pat, text[:5000], re.IGNORECASE)
        if m:
            return f"{m.group(1)}+ yrs"
    return ""


def extract_notice_period(text: str) -> str:
    patterns = [
        r"(?i)notice\s*period\s*[:\-]?\s*([^\n.]{3,40})",
        r"(?i)(immediate\s*joiner|immediate\s*availability|available\s*immediately)",
        r"(?i)(serving\s*notice|currently\s*serving)",
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            val = m.group(1).strip()[:50]
            return val
    return ""


def extract_skills(text: str) -> str:
    # Broad skill list relevant to Indian IT market
    skills_db = [
        # Languages
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go", "Rust",
        "Ruby", "PHP", "Swift", "Kotlin", "Scala", "R", "MATLAB", "Perl", "Shell",
        # Web
        "React", "Angular", "Vue", "Node.js", "Django", "Flask", "FastAPI",
        "Spring Boot", "Spring", "HTML", "CSS", "GraphQL", "REST API",
        # Cloud
        "AWS", "Azure", "GCP", "Google Cloud", "Oracle Cloud",
        # Data / ML
        "Machine Learning", "Deep Learning", "NLP", "TensorFlow", "PyTorch",
        "Pandas", "NumPy", "Scikit-learn", "Spark", "PySpark", "Hadoop",
        "Power BI", "Tableau", "Databricks", "Snowflake", "dbt",
        "Data Engineering", "Data Science", "Data Analysis",
        # Databases
        "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Oracle", "Cassandra",
        "DynamoDB", "Cosmos DB", "SQL Server",
        # DevOps / Infra
        "Docker", "Kubernetes", "CI/CD", "Jenkins", "Terraform", "Ansible",
        "Git", "GitHub", "GitLab", "Linux", "Unix",
        # Other
        "Agile", "Scrum", "Microservices", "SAP", "Salesforce", "ServiceNow",
        "Selenium", "Postman", "JIRA",
    ]
    lower = text.lower()
    found = [s for s in skills_db if s.lower() in lower]
    # Deduplicate case-insensitively
    seen, unique = set(), []
    for s in found:
        if s.lower() not in seen:
            seen.add(s.lower())
            unique.append(s)
    return ", ".join(unique[:18])


def extract_education(text: str) -> str:
    patterns = [
        r"(?i)(ph\.?d\.?|doctorate)",
        r"(?i)(m\.?tech|m\.?e\.?|master\s+of\s+(?:technology|engineering|science|computer))",
        r"(?i)(mba|master\s+of\s+business)",
        r"(?i)(m\.?sc\.?|master\s+of\s+science)",
        r"(?i)(b\.?tech|b\.?e\.?|bachelor\s+of\s+(?:technology|engineering))",
        r"(?i)(b\.?sc\.?|b\.?c\.?a\.?|b\.?com|bachelor)",
        r"(?i)(diploma)",
    ]
    for pat in patterns:
        if re.search(pat, text):
            m = re.search(pat, text)
            return m.group(1).strip().upper()
    return ""


def get_download_date(path: Path) -> str:
    stat = path.stat()
    # st_birthtime = creation time on macOS (= download date)
    ts = getattr(stat, "st_birthtime", stat.st_mtime)
    return datetime.fromtimestamp(ts).strftime("%Y-%m-%d")


# ── Main processing ────────────────────────────────────────────────────────────

def process_file(path: Path) -> dict | None:
    text = extract_text(path)

    if not is_likely_resume(text, path.name):
        return None

    emails = extract_emails(text)
    phones = extract_phones(text)

    return {
        "Document Name":    path.name,
        "Candidate Name":   extract_name(text, path.name),
        "Current Role":     extract_current_role(text, path.name),
        "Email":            emails[0] if emails else "",
        "Phone":            phones[0] if phones else "",
        "LinkedIn":         extract_linkedin(text),
        "Location":         extract_location(text),
        "Experience":       extract_experience_years(text),
        "Notice Period":    extract_notice_period(text),
        "Education":        extract_education(text),
        "Key Skills":       extract_skills(text),
        "Downloaded Date":  get_download_date(path),
        "File Type":        path.suffix.upper().lstrip("."),
        "Status":           "New",          # pipeline stage — edit manually
        "Interview Round":  "",             # fill as you progress
        "Rating":           "",             # 1-5 star rating slot
        "Notes":            "",             # free-text notes column
    }


# ── Excel generation ───────────────────────────────────────────────────────────

# Column widths (chars)
COL_WIDTHS = {
    "Document Name":   36,
    "Candidate Name":  22,
    "Current Role":    28,
    "Email":           28,
    "Phone":           18,
    "LinkedIn":        30,
    "Location":        18,
    "Experience":      13,
    "Notice Period":   16,
    "Education":       14,
    "Key Skills":      48,
    "Downloaded Date": 16,
    "File Type":       10,
    "Status":          16,
    "Interview Round": 16,
    "Rating":          10,
    "Notes":           30,
}

HEADER_BG   = "1F4E79"   # dark navy
ALT_ROW_BG  = "DCE6F1"   # light blue
STATUS_COLORS = {         # conditional colour for Status column
    "New":          "FFFF99",
    "Shortlisted":  "C6EFCE",
    "Interviewed":  "FFEB9C",
    "Selected":     "A9D18E",
    "Rejected":     "FFC7CE",
    "On Hold":      "D9D9D9",
}


def _border(color="BFBFBF"):
    s = Side(style="thin", color=color)
    return Border(left=s, right=s, top=s, bottom=s)


def create_excel(records: list[dict], out_path: Path):
    wb = openpyxl.Workbook()

    # ── Candidates sheet ─────────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Candidates"

    headers = list(COL_WIDTHS.keys())

    # Header row
    hdr_font  = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
    hdr_fill  = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font      = hdr_font
        cell.fill      = hdr_fill
        cell.alignment = hdr_align
        cell.border    = _border("FFFFFF")
        ws.column_dimensions[get_column_letter(c)].width = COL_WIDTHS[h]

    ws.row_dimensions[1].height = 28

    # Data rows
    data_font = Font(name="Calibri", size=10)
    alt_fill  = PatternFill(start_color=ALT_ROW_BG, end_color=ALT_ROW_BG, fill_type="solid")
    status_col = headers.index("Status") + 1

    for r, rec in enumerate(records, 2):
        bg = alt_fill if r % 2 == 0 else PatternFill()
        for c, h in enumerate(headers, 1):
            val  = rec.get(h, "")
            cell = ws.cell(row=r, column=c, value=val)
            cell.font      = data_font
            cell.alignment = Alignment(vertical="center", wrap_text=False)
            cell.border    = _border()
            if h == "Status" and val in STATUS_COLORS:
                cell.fill = PatternFill(
                    start_color=STATUS_COLORS[val],
                    end_color=STATUS_COLORS[val],
                    fill_type="solid"
                )
            elif bg.fill_type:
                cell.fill = bg
        ws.row_dimensions[r].height = 18

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    # ── Summary sheet ─────────────────────────────────────────────────────────
    ws2 = wb.create_sheet("Summary")
    ws2.column_dimensions["A"].width = 28
    ws2.column_dimensions["B"].width = 20

    def _s(row, label, value, bold_val=False):
        c1 = ws2.cell(row=row, column=1, value=label)
        c2 = ws2.cell(row=row, column=2, value=value)
        c1.font = Font(bold=True, name="Calibri", size=11)
        c2.font = Font(bold=bold_val, name="Calibri", size=11,
                       color="1F4E79" if bold_val else "000000")
        c2.alignment = Alignment(horizontal="left")

    title = ws2.cell(row=1, column=1, value="Resume Scanner — TechyEra")
    title.font = Font(bold=True, size=14, color="1F4E79", name="Calibri")
    ws2.merge_cells("A1:B1")

    _s(3,  "Total resumes scanned:",     len(records), True)
    _s(4,  "With email address:",        sum(1 for r in records if r["Email"]))
    _s(5,  "With phone number:",         sum(1 for r in records if r["Phone"]))
    _s(6,  "With LinkedIn profile:",     sum(1 for r in records if r["LinkedIn"]))
    _s(7,  "With experience info:",      sum(1 for r in records if r["Experience"]))
    _s(8,  "With notice period info:",   sum(1 for r in records if r["Notice Period"]))
    _s(9,  "Scanned on:",                datetime.now().strftime("%Y-%m-%d %H:%M"))
    _s(10, "Source folder:",             str(DOWNLOADS))

    # Status legend
    ws2.cell(row=12, column=1, value="Status colour legend:").font = Font(bold=True, name="Calibri")
    for i, (status, color) in enumerate(STATUS_COLORS.items(), 13):
        cell = ws2.cell(row=i, column=1, value=status)
        cell.fill   = PatternFill(start_color=color, end_color=color, fill_type="solid")
        cell.font   = Font(name="Calibri", size=10)
        cell.border = _border()

    wb.save(str(out_path))


# ── Entry point ────────────────────────────────────────────────────────────────

def main():
    print("=" * 62)
    print("  TechyEra Resume Scanner")
    print("=" * 62)
    print(f"  Folder : {DOWNLOADS}")
    print()

    # Collect candidate files (skip temp/lock files upfront)
    candidates = [
        f for f in sorted(DOWNLOADS.iterdir())
        if f.is_file()
        and f.suffix.lower() in RESUME_EXTENSIONS
        and not f.name.startswith(".~")
        and not f.name.startswith("~$")
    ]
    print(f"  Files with resume extensions found : {len(candidates)}")

    records, skipped, errors = [], [], []

    for path in candidates:
        print(f"  ▶ {path.name[:70]}")
        try:
            rec = process_file(path)
            if rec:
                records.append(rec)
                email_tag = f"✉ {rec['Email']}" if rec["Email"] else "no email"
                phone_tag = f"☎ {rec['Phone']}" if rec["Phone"] else "no phone"
                print(f"       → {rec['Candidate Name']} | {email_tag} | {phone_tag}")
            else:
                skipped.append(path.name)
                print(f"       → skipped (not a resume)")
        except Exception as e:
            errors.append((path.name, str(e)))
            print(f"       → ERROR: {e}")

    # Save Excel
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = DOWNLOADS / f"TechyEra_Candidates_{ts}.xlsx"
    create_excel(records, out_path)

    print()
    print("=" * 62)
    print(f"  ✅  Resumes extracted : {len(records)}")
    print(f"  ⏭   Skipped (not resume) : {len(skipped)}")
    if errors:
        print(f"  ❌  Errors           : {len(errors)}")
        for name, err in errors:
            print(f"       {name}: {err}")
    print(f"  📄  Output saved to  : {out_path.name}")
    print("=" * 62)

    # Open in Excel / Numbers
    subprocess.run(["open", str(out_path)], check=False)


if __name__ == "__main__":
    main()
