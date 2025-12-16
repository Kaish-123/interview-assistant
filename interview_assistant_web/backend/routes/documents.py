"""
Document Routes - Handle resume/JD uploads
"""
import base64
import io
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session

from database.db import get_db, Document
from models.schemas import DocumentCreate, DocumentResponse
from services.chat_service import chat_service

router = APIRouter(prefix="/documents", tags=["Documents"])


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from various document formats"""
    extension = filename.lower().split('.')[-1] if '.' in filename else ''
    
    try:
        if extension == 'txt':
            return file_content.decode('utf-8')
        
        elif extension == 'pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_content))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text.strip()
        
        elif extension == 'docx':
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip()
        
        elif extension == 'doc':
            # Older .doc format - try textract if available
            try:
                import textract
                return textract.process(io.BytesIO(file_content)).decode('utf-8')
            except ImportError:
                raise HTTPException(
                    status_code=400,
                    detail=".doc format requires textract library"
                )
        
        else:
            # Try to read as plain text
            try:
                return file_content.decode('utf-8')
            except:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file format: {extension}"
                )
                
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to extract text: {str(e)}"
        )


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    doc_type: str = Form(default="resume"),
    session_id: int = Form(default=None),
    db: Session = Depends(get_db)
):
    """Upload and process a document (resume, JD, etc.)"""
    
    # Read file content
    file_content = await file.read()
    
    # Extract text
    text = extract_text_from_file(file_content, file.filename)
    
    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from document")
    
    # Create document record
    doc = Document(
        filename=file.filename,
        doc_type=doc_type,
        content=text,
        session_id=session_id
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    # If session_id provided, add as system context
    if session_id:
        chat_service.add_document_context(
            db,
            session_id=session_id,
            filename=file.filename,
            content=text,
            doc_type=doc_type
        )
    
    return {
        "success": True,
        "document": DocumentResponse(
            id=doc.id,
            filename=doc.filename,
            doc_type=doc.doc_type,
            content=text[:500] + "..." if len(text) > 500 else text,  # Preview
            session_id=doc.session_id,
            created_at=doc.created_at
        ),
        "text_length": len(text)
    }


@router.get("/", response_model=List[DocumentResponse])
def get_documents(
    doc_type: str = None,
    session_id: int = None,
    db: Session = Depends(get_db)
):
    """Get all documents, optionally filtered"""
    query = db.query(Document)
    
    if doc_type:
        query = query.filter(Document.doc_type == doc_type)
    if session_id:
        query = query.filter(Document.session_id == session_id)
    
    documents = query.order_by(Document.created_at.desc()).all()
    
    return [
        DocumentResponse(
            id=d.id,
            filename=d.filename,
            doc_type=d.doc_type,
            content=d.content[:500] + "..." if len(d.content) > 500 else d.content,
            session_id=d.session_id,
            created_at=d.created_at
        )
        for d in documents
    ]


@router.get("/{document_id}", response_model=DocumentResponse)
def get_document(document_id: int, db: Session = Depends(get_db)):
    """Get a specific document"""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return DocumentResponse(
        id=doc.id,
        filename=doc.filename,
        doc_type=doc.doc_type,
        content=doc.content,
        session_id=doc.session_id,
        created_at=doc.created_at
    )


@router.delete("/{document_id}")
def delete_document(document_id: int, db: Session = Depends(get_db)):
    """Delete a document"""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    db.delete(doc)
    db.commit()
    return {"success": True}


@router.post("/{document_id}/add-to-session")
def add_document_to_session(
    document_id: int,
    session_id: int,
    db: Session = Depends(get_db)
):
    """Add existing document to a chat session"""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    session = chat_service.get_session(db, session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    chat_service.add_document_context(
        db,
        session_id=session_id,
        filename=doc.filename,
        content=doc.content,
        doc_type=doc.doc_type
    )
    
    return {"success": True, "message": f"Added {doc.filename} to session"}





